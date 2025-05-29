#!/usr/bin/env python3
"""
Script to create a Word template for the dossier generator.
This creates a basic template that can be used with docxtpl.
"""

from docxtpl import DocxTemplate
import tempfile
import os

def create_dossier_template():
    """Create a Word template for evidence dossiers"""
    
    # Create a temporary document to serve as base
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        # Create a basic document with python-docx first
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Evidence Dossier', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Case information section
        doc.add_heading('Case Information', level=1)
        
        case_info = doc.add_paragraph()
        case_info.add_run('Case ID: ').bold = True
        case_info.add_run('{{case_id}}')
        
        generation_info = doc.add_paragraph()
        generation_info.add_run('Generated: ').bold = True
        generation_info.add_run('{{generation_date}}')
        
        total_info = doc.add_paragraph()
        total_info.add_run('Total Evidence Items: ').bold = True
        total_info.add_run('{{total_items}}')
        
        # Evidence summary section
        doc.add_heading('Evidence Summary', level=1)
        
        # Template for evidence items loop
        doc.add_paragraph('{% for evidence in evidence_items %}')
        
        # Evidence item heading
        evidence_heading = doc.add_heading('Evidence Item {{loop.index}}', level=2)
        
        # Evidence details table template
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Table headers and template values
        cells = table.rows[0].cells
        cells[0].text = 'Filename:'
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = '{{evidence.filename}}'
        
        cells = table.rows[1].cells
        cells[0].text = 'SHA-256 Hash:'
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = '{{evidence.sha256}}'
        
        cells = table.rows[2].cells
        cells[0].text = 'Captured Date:'
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = '{{evidence.captured_at}}'
        
        cells = table.rows[3].cells
        cells[0].text = 'File Type:'
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = '{{evidence.mime_type}}'
        
        cells = table.rows[4].cells
        cells[0].text = 'File Size:'
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = '{{evidence.size_bytes}} bytes'
        
        cells = table.rows[5].cells
        cells[0].text = 'GPS Coordinates:'
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = '{% if evidence.coordinates %}{{evidence.coordinates.latitude}}, {{evidence.coordinates.longitude}}{% else %}N/A{% endif %}'
        
        # End of loop
        doc.add_paragraph('{% endfor %}')
        
        # Chain of custody section
        doc.add_heading('Chain of Custody', level=1)
        
        custody_p1 = doc.add_paragraph(
            'This dossier was generated from the Evidence-MVP system. All files included in this report have been '
            'cryptographically verified and stored with tamper-evident controls.'
        )
        
        custody_p2 = doc.add_paragraph(
            'Digital signatures and hash verification ensure the integrity of all evidence items.'
        )
        
        # Save the document
        doc.save(tmp.name)
          # Move to final location
        template_dir = os.path.join(os.path.dirname(__file__), 'templates')
        os.makedirs(template_dir, exist_ok=True)
        template_path = os.path.join(template_dir, 'dossier_template.docx')
        
        import shutil
        shutil.move(tmp.name, template_path)
        
        print(f"Template created at: {template_path}")
        return template_path

if __name__ == "__main__":
    create_dossier_template()
